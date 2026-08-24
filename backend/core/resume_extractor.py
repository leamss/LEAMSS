"""Phase 6.7 Part 2 — Resume Upload + AI Extraction.

Extracts structured candidate profile fields from a resume (PDF or DOCX) using:
  1. Text extraction via pdfplumber / python-docx
  2. Claude AI (Sonnet 4.6) for structured JSON extraction

Output schema matches the Phase 6.7 ProfileCreate model so it can be used to
prefill the wizard directly.
"""
import asyncio
import io
import json
import logging
import os
from typing import Dict, Any, List, Optional, Tuple
from openai import AsyncOpenAI
import re


logger = logging.getLogger(__name__)

PERPLEXITY_API_KEY = os.environ.get("PERPLEXITY_API_KEY", "")
PERPLEXITY_MODEL = "sonar-pro"

MAX_TEXT_CHARS = 24_000  # safe Claude input budget


EXTRACTION_PROMPT = """You are a resume-parsing assistant for an immigration consultancy.

Extract the candidate's profile from the resume text below. Return ONLY a JSON object
matching this exact schema (omit fields if not found — do NOT invent data):

{
  "name": "Full name of the candidate",
  "email": "email if present",
  "phone": "phone with country code if present",
  "marital_status": "single | married | de_facto | separated | divorced | widowed (if mentioned)",
  "primary_applicant": {
    "personal": {
      "full_name": "...",
      "date_of_birth": "YYYY-MM-DD if found",
      "age": null,
      "gender": "male | female | other (if found)",
      "nationality": "if found",
      "current_country": "country candidate currently lives in",
      "current_city": "city candidate currently lives in"
    },
    "professional": {
      "current_profession": "MOST RECENT job role (e.g., 'Marketing Specialist', 'Software Engineer')",
      "designation": "Most recent designation/title",
      "years_experience_total": 0.0,
      "years_in_current_role": 0.0,
      "industry": "Industry sector",
      "employer_name": "Current employer",
      "salary_inr_per_annum": null,
      "has_managerial_experience": false
    },
    "education": {
      "highest_qualification": "doctorate | master | bachelor | diploma | trade | high_school",
      "field_of_study": "Field of study (e.g., 'Computer Science')",
      "institution": "Last/highest institution",
      "country": "Country of highest qualification",
      "year_completed": null
    },
    "language": {
      "primary_test": "IELTS | PTE | TOEFL | none",
      "test_completed": false,
      "test_date": null,
      "scores": {}
    },
    "work_history": [
      {"employer": "...", "designation": "...", "start_date": "YYYY-MM", "end_date": "YYYY-MM or null", "country": "...", "duties": "1-2 line summary"}
    ]
  },
  "_extraction_notes": ["any caveats, ambiguities, or fields that need user review"]
}

CRITICAL RULES:
- The MOST RECENT job (top of work history) = `current_profession` and `designation`.
- `years_experience_total` = sum of years across all professional roles (not student/intern).
- `highest_qualification` MUST be one of: doctorate | master | bachelor | diploma | trade | high_school. Map common synonyms (PhD→doctorate, M.Tech/MBA/M.S→master, B.Tech/B.E/B.S→bachelor).
- DO NOT confuse education field with current profession. e.g., A B.V.Sc graduate currently working as Marketing Specialist → `current_profession='Marketing Specialist'`, `field_of_study='Veterinary Science'`.
- If language test is mentioned (IELTS/PTE/TOEFL), set `test_completed=true` and include scores. Otherwise default to `primary_test='IELTS', test_completed=false, scores={}`.
- For unclear fields, leave them null/empty and add a note in `_extraction_notes`.
- Output MUST be valid JSON. NO markdown, NO commentary outside the JSON.
"""


def extract_text_from_pdf(file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
    """Extract plain text from a PDF resume."""
    import pdfplumber

    text_parts = []
    page_count = 0
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            t = page.extract_text() or ""
            if t.strip():
                text_parts.append(t)
    text = "\n\n".join(text_parts)
    return text, {"page_count": page_count, "char_count": len(text)}


def extract_text_from_docx(file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
    """Extract plain text from a DOCX resume."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # Also pull table cells
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text.strip())
    text = "\n".join(parts)
    return text, {"paragraph_count": len(doc.paragraphs), "char_count": len(text)}


def extract_text(filename: str, file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
    """Dispatch by file extension."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    if name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    if name.endswith(".txt"):
        try:
            t = file_bytes.decode("utf-8", errors="replace")
        except Exception:
            t = file_bytes.decode("latin-1", errors="replace")
        return t, {"char_count": len(t)}
    raise ValueError(f"Unsupported file type — only .pdf, .docx, .txt allowed (got '{name}')")


# ---------------------------------------------------------------------------
# OCR fallback (scanned/image resumes) via vision LLM
# ---------------------------------------------------------------------------
OCR_MODEL = "gemini-2.5-flash"
OCR_MAX_PAGES = 4
_OCR_SYS = ("You are an OCR engine. Transcribe ALL text from the document image(s) verbatim "
            "as plain text. Preserve names, job titles, companies, dates, skills and education. "
            "Output ONLY the transcribed text, no commentary.")


async def ocr_images_to_text(images_b64: List[str], model: Optional[str] = None) -> str:
    """Transcribe text from one or more images using vision LLM. Returns plain text ('' on failure)."""
    return ""


def _render_pdf_to_pngs(file_bytes: bytes, max_pages: int = OCR_MAX_PAGES, dpi: int = 150) -> List[str]:
    """Render the first N PDF pages to base64 PNGs (CPU-bound; call via to_thread)."""
    import base64
    import pymupdf
    out: List[str] = []
    doc = pymupdf.open(stream=file_bytes, filetype="pdf")
    try:
        for i in range(min(max_pages, doc.page_count)):
            pix = doc[i].get_pixmap(dpi=dpi)
            out.append(base64.b64encode(pix.tobytes("png")).decode())
    finally:
        doc.close()
    return out


async def ocr_pdf_bytes(file_bytes: bytes, model: Optional[str] = None) -> str:
    try:
        pngs = await asyncio.to_thread(_render_pdf_to_pngs, file_bytes)
    except Exception as e:  # noqa: BLE001
        logger.warning(f"PDF render for OCR failed: {e}")
        return ""
    return await ocr_images_to_text(pngs, model=model) if pngs else ""


async def ocr_image_bytes(file_bytes: bytes, model: Optional[str] = None) -> str:
    import base64
    return await ocr_images_to_text([base64.b64encode(file_bytes).decode()], model=model)


async def extract_text_smart(filename: str, file_bytes: bytes) -> Tuple[Optional[str], Optional[str]]:
    """Extract resume text with an OCR fallback for scanned PDFs & image files.
    Returns (text, error_message)."""
    name = (filename or "").lower()
    is_image = (name.endswith((".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp", ".tif", ".tiff"))
                or file_bytes[:3] == b"\xff\xd8\xff" or file_bytes[:8] == b"\x89PNG\r\n\x1a\n")
    is_pdf = name.endswith(".pdf") or file_bytes[:4] == b"%PDF"

    if is_image:
        text = await ocr_image_bytes(file_bytes)
        if text and len(text.strip()) >= 30:
            return text, None
        return None, "Could not read text from image."

    try:
        text, _meta = await asyncio.to_thread(extract_text, filename, file_bytes)
    except ValueError as e:
        if is_pdf:
            ocr = await ocr_pdf_bytes(file_bytes)
            if ocr and len(ocr.strip()) >= 30:
                return ocr, None
        return None, str(e)

    if (not text or len(text.strip()) < 50) and is_pdf:
        ocr = await ocr_pdf_bytes(file_bytes)  # scanned/image-only PDF
        if ocr and len(ocr.strip()) >= 30:
            return ocr, None
    if not text or len(text.strip()) < 30:
        return None, "Resume text was empty (scanned image or unreadable file)."
    return text, None


async def parse_resume_with_ai(
    resume_text: str,
    session_id: Optional[str] = None,
    model: Optional[str] = None,
    **kwargs: Any,
) -> Dict[str, Any]:
    """Send resume text to AI for structured extraction.
    Returns the parsed JSON (Phase 6.7 ProfileCreate shape) or a fallback empty shell.
    """
    key = os.environ.get("PERPLEXITY_API_KEY", "").strip()
    if not key:
        return {"_error": "PERPLEXITY_API_KEY not configured"}
    if not resume_text or len(resume_text.strip()) < 50:
        return {"_error": "Resume text is too short to extract anything meaningful"}

    # Trim to safe budget
    text = resume_text[:MAX_TEXT_CHARS]

    # try:
    #     from emergentintegrations.llm.chat import LlmChat, UserMessage  # type: ignore
    # except ImportError as e:
    #     return {"_error": f"emergentintegrations missing: {e}"}

    sid = session_id or f"resume-{os.urandom(4).hex()}"
    user_prompt = (
        "## RESUME TEXT\n```\n"
        + text
        + "\n```\n\nReturn the extracted JSON now."
    )

    try:
        import httpx as _httpx
        _http = _httpx.AsyncClient(verify=False, timeout=60)
        client = AsyncOpenAI(
            api_key=key,
            base_url="https://api.perplexity.ai",
            http_client=_http,
        )

        response = None
        for attempt in range(5):
            try:
                response = await client.chat.completions.create(
                    model=model or PERPLEXITY_MODEL,
                    temperature=0,
                    max_tokens=2500,
                    messages=[
                        {
                            "role": "system",
                            "content": EXTRACTION_PROMPT,
                        },
                        {
                            "role": "user",
                            "content": user_prompt,
                        },
                    ],
                )
                break
            except Exception as re_err:
                err_str = str(re_err).lower()
                if ("429" in err_str or "rate limit" in err_str) and attempt < 4:
                    wait_sec = 2.0 * (attempt + 1) + 0.5
                    logger.warning(f"Perplexity rate limited on resume extraction, retrying in {wait_sec:.1f}s...")
                    await asyncio.sleep(wait_sec)
                    continue
                raise

        if not response:
            return {"_error": "No response from AI"}

        message = response.choices[0].message
        raw = message.content or ""
        logger.info("Perplexity resume extraction response received (length=%d)", len(raw))
        

      

        if raw.startswith("```"):
            raw = raw.strip("`").replace("json", "", 1).strip()



                # Remove markdown fences if present
        raw = re.sub(
            r"^```(?:json)?",
            "",
            raw.strip(),
            flags=re.IGNORECASE
        )
        raw = re.sub(
            r"```$",
            "",
            raw.strip()
        )

        # Find the first JSON object
               # Remove markdown fences if present
        

        # Find the JSON object
        match = re.search(r"\{.*\}", raw, re.DOTALL)

        if not match:
            logger.warning("AI returned non-JSON response: %s", raw[:300])
            return {
                "_error": "AI returned non-JSON response",
                "_raw": raw[:1000],
            }

        json_text = match.group(0)

        # Remove control characters
        json_text = re.sub(r"[\x00-\x1F\x7F]", "", json_text)

        # Remove trailing commas
        json_text = re.sub(r",(\s*[}\]])", r"\1", json_text)

        try:
            parsed = json.loads(json_text)

        except json.JSONDecodeError as e:
            logger.warning("Resume JSON decode error: %s", e)
            start = max(0, e.pos - 300)
            end = min(len(json_text), e.pos + 300)
            return {
                "_error": f"Invalid JSON: {e}",
                "_raw": json_text[start:end]
            }

        parsed["_ai_status"] = "ok"
        parsed["_ai_model"] = PERPLEXITY_MODEL

        return parsed

        
        

    except json.JSONDecodeError as e:
        logger.error(f"Resume JSON error: {e}")
        return {
            "_error": f"Invalid JSON: {e}"
        }

    except Exception as e:
        logger.exception("Resume parsing failed")
        return {
            "_error": f"Perplexity API failed: {type(e).__name__}: {str(e)}"
        }